from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor


TEMPLATE_PATH = Path("templates/cv_recruiter_template.docx")
ACCENT_COLOR = RGBColor(44, 82, 100)
MUTED_COLOR = RGBColor(90, 90, 90)


def main() -> int:
    create_template(TEMPLATE_PATH)
    print(TEMPLATE_PATH)
    return 0


def create_template(path: str | Path = TEMPLATE_PATH) -> Path:
    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _set_style(document)
    _add_header(document)
    _add_section(document, "Profil")
    document.add_paragraph("{{ proposed_title }}")
    document.add_paragraph("{{ targeted_summary }}")
    _add_section(document, "Compétences clés")
    document.add_paragraph("{% for skill in skills %}{{ skill }}{% if not loop.last %} • {% endif %}{% endfor %}")
    _add_section(document, "Expériences")
    document.add_paragraph("{% for experience in experiences %}")
    document.add_paragraph("{{ experience.heading }}")
    document.add_paragraph("{% for bullet in experience.bullets %}")
    document.add_paragraph("{{ bullet }}", style="List Bullet")
    document.add_paragraph("{% endfor %}")
    document.add_paragraph("{% endfor %}")
    _add_section(document, "Projets")
    document.add_paragraph("{% for project in projects %}")
    document.add_paragraph("{{ project.name }}")
    document.add_paragraph("{% if project.url %}{{ project.url }}{% endif %}")
    document.add_paragraph("{% for bullet in project.bullets %}")
    document.add_paragraph("{{ bullet }}", style="List Bullet")
    document.add_paragraph("{% endfor %}")
    document.add_paragraph("{% if project.technologies %}Technologies : {{ project.technologies }}{% endif %}")
    document.add_paragraph("{% endfor %}")
    _add_section(document, "Formation")
    document.add_paragraph("{% for item in education %}")
    document.add_paragraph("{{ item }}", style="List Bullet")
    document.add_paragraph("{% endfor %}")

    document.save(output_path)
    return output_path


def _set_style(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Cm(1.35)
        section.bottom_margin = Cm(1.35)
        section.left_margin = Cm(1.45)
        section.right_margin = Cm(1.45)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    bullet = document.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(10.5)
    bullet.paragraph_format.space_after = Pt(1)


def _add_header(document: Document) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run("{{ identity.name }}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = ACCENT_COLOR

    paragraph = document.add_paragraph()
    run = paragraph.add_run(
        "{{ identity.title }}{% if identity.location %} | {{ identity.location }}{% endif %}"
        "{% if identity.email %} | {{ identity.email }}{% endif %}"
        "{% if identity.phone %} | {{ identity.phone }}{% endif %}"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED_COLOR

    paragraph = document.add_paragraph()
    run = paragraph.add_run("CV - {{ identity.name or proposed_title }}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = ACCENT_COLOR

    paragraph = document.add_paragraph()
    run = paragraph.add_run("─" * 58)
    run.font.size = Pt(5)
    run.font.color.rgb = RGBColor(190, 196, 200)


def _add_section(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(12.5)
    run.font.color.rgb = ACCENT_COLOR


if __name__ == "__main__":
    raise SystemExit(main())
