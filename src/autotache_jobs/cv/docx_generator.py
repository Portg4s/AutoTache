"""DOCX generation for local targeted CV drafts."""

from __future__ import annotations

import re
import unicodedata
from pathlib import Path
from typing import Any, Literal

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor

from autotache_jobs.cv.builder import CvExperience, CvProject, TargetedCvData, build_targeted_cv_data
from autotache_jobs.cv.profile import CvProfile


CvRenderMode = Literal["draft", "recruiter"]
ACCENT_COLOR = RGBColor(44, 82, 100)
MUTED_COLOR = RGBColor(90, 90, 90)


def generate_cv_docx(
    *,
    offer: dict[str, Any],
    profile: CvProfile,
    output_dir: str | Path = "generated_cv",
    mode: CvRenderMode = "draft",
) -> Path:
    _validate_mode(mode)
    target_dir = Path(output_dir)
    target_dir.mkdir(parents=True, exist_ok=True)

    cv_data = build_targeted_cv_data(offer, profile)
    output_path = target_dir / _safe_filename(cv_data, mode)
    document = _build_document(cv_data, mode=mode)
    document.save(output_path)
    return output_path


def _build_document(cv_data: TargetedCvData, *, mode: CvRenderMode) -> Document:
    document = Document()
    _set_base_style(document)

    if mode == "recruiter":
        _add_recruiter_content(document, cv_data)
    else:
        _add_draft_content(document, cv_data)

    return document


def _add_draft_content(document: Document, cv_data: TargetedCvData) -> None:
    offer = cv_data.offer_info
    _add_candidate_header(document, cv_data)
    document.add_paragraph(f"CV ciblé — {offer.company} — {offer.title}", style="Title")
    document.add_paragraph(
        " | ".join([offer.location, offer.contract_type, offer.source, f"Score {offer.score}"])
    )

    document.add_heading("Profil ciblé", level=1)
    document.add_heading("Titre proposé", level=2)
    document.add_paragraph(cv_data.proposed_title)
    document.add_heading("Accroche ciblée", level=2)
    document.add_paragraph(cv_data.targeted_summary)

    document.add_heading("Compétences clés", level=1)
    _add_skills(document, cv_data, include_to_confirm=True)

    document.add_heading("Expériences", level=1)
    _add_experiences(document, cv_data.experiences)

    document.add_heading("Projets", level=1)
    _add_projects(document, cv_data.projects, include_source=True)

    document.add_heading("Formation", level=1)
    _add_bullets(document, cv_data.education, empty="Aucune formation structurée détectée dans le profil maître.")

    document.add_heading("Informations offre", level=1)
    _add_bullets(
        document,
        [
            f"Titre de l'offre: {offer.title}",
            f"Entreprise: {offer.company}",
            f"Source: {offer.source}",
            f"Localisation: {offer.location}",
            f"Type de contrat: {offer.contract_type}",
            f"Decision AutoTache: {offer.decision}",
            f"Score: {offer.score}",
            f"URL: {offer.url}",
        ],
    )

    document.add_heading("Analyse de correspondance", level=1)
    document.add_heading("Correspondance profil/offre", level=2)
    _add_bullets(
        document,
        [
            "Compétences fortes présentes: " + _inline_list(cv_data.skills.confirmed),
            "Compétences moyennes présentes: " + _inline_list(cv_data.skills.complementary),
            "Compétences à confirmer présentes: " + _inline_list(cv_data.skills.to_confirm),
        ],
    )
    document.add_heading("Mots-clés détectés dans l’offre", level=2)
    _add_bullets(document, cv_data.analysis.keywords, empty="Aucun mot-clé utile hors profil détecté.")

    document.add_heading("Points de prudence", level=1)
    document.add_heading("Points à ne pas sur-vendre", level=2)
    _add_bullets(
        document,
        cv_data.analysis.oversell_points,
        empty="Aucune compétence à confirmer détectée dans l'offre.",
    )
    document.add_heading("Règles de prudence", level=2)
    _add_bullets(document, cv_data.analysis.caution_rules)


def _add_recruiter_content(document: Document, cv_data: TargetedCvData) -> None:
    _add_recruiter_header(document, cv_data)
    title = cv_data.identity.name or cv_data.proposed_title
    _add_recruiter_title(document, f"CV - {title}")

    _add_recruiter_section_heading(document, "Profil")
    _add_profile_title(document, cv_data.proposed_title)
    _add_profile_summary(document, cv_data.recruiter_summary)

    _add_recruiter_section_heading(document, "Compétences clés")
    _add_recruiter_skills(document, cv_data)

    _add_recruiter_section_heading(document, "Expériences")
    _add_recruiter_experiences(document, cv_data.experiences)

    _add_recruiter_section_heading(document, "Projets")
    _add_recruiter_projects(document, cv_data.projects)

    _add_recruiter_section_heading(document, "Formation")
    _add_bullets(document, cv_data.education, empty="Aucune formation structurée détectée dans le profil maître.")


def _set_base_style(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Cm(1.35)
        section.bottom_margin = Cm(1.35)
        section.left_margin = Cm(1.45)
        section.right_margin = Cm(1.45)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    title = document.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(16)
    title.paragraph_format.space_before = Pt(2)
    title.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(3)

    bullet = document.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(10.5)
    bullet.paragraph_format.space_after = Pt(1)


def _add_candidate_header(document: Document, cv_data: TargetedCvData) -> None:
    identity = cv_data.identity
    if identity.name:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(identity.name)
        run.bold = True
        run.font.size = Pt(13)

    details = [identity.title, identity.location, identity.email, identity.phone]
    details = [detail for detail in details if detail]
    if details:
        document.add_paragraph(" | ".join(details))


def _add_recruiter_header(document: Document, cv_data: TargetedCvData) -> None:
    identity = cv_data.identity
    if identity.name:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(identity.name)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = ACCENT_COLOR

    details = [identity.title, identity.location, identity.email, identity.phone]
    details = [detail for detail in details if detail]
    if details:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(5)
        run = paragraph.add_run(" | ".join(details))
        run.font.size = Pt(10)
        run.font.color.rgb = MUTED_COLOR


def _add_recruiter_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Title")
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = ACCENT_COLOR
    _add_separator(document)


def _add_separator(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run("─" * 58)
    run.font.size = Pt(5)
    run.font.color.rgb = RGBColor(190, 196, 200)


def _add_recruiter_section_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12.5)
    run.font.color.rgb = ACCENT_COLOR


def _add_profile_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)


def _add_profile_summary(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.08
    paragraph.add_run(text)


def _add_skills(document: Document, cv_data: TargetedCvData, *, include_to_confirm: bool) -> None:
    if not include_to_confirm:
        _add_bullets(document, cv_data.skills.confirmed + cv_data.skills.complementary)
        return

    if cv_data.skills.confirmed or include_to_confirm:
        document.add_heading("Compétences confirmées pertinentes", level=2)
        _add_bullets(
            document,
            cv_data.skills.confirmed,
            empty="Aucune compétence forte du profil détectée explicitement dans l'offre.",
        )
    if cv_data.skills.complementary or include_to_confirm:
        document.add_heading("Compétences complémentaires", level=2)
        _add_bullets(
            document,
            cv_data.skills.complementary,
            empty="Aucune compétence complémentaire du profil détectée explicitement dans l'offre.",
        )
    if not include_to_confirm:
        return

    document.add_heading("Compétences à confirmer", level=2)
    if cv_data.skills.to_confirm:
        _add_bullets(
            document,
            [
                f"{skill} (à confirmer, ne pas présenter comme maîtrisé)"
                for skill in cv_data.skills.to_confirm
            ],
        )
    else:
        _add_bullets(document, [], empty="Aucune compétence à confirmer détectée dans l'offre.")


def _add_recruiter_skills(document: Document, cv_data: TargetedCvData) -> None:
    skills = cv_data.skills.confirmed + cv_data.skills.complementary
    if not skills:
        return

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(" • ".join(skills))
    run.font.size = Pt(10.5)


def _add_recruiter_experiences(document: Document, experiences: list[CvExperience]) -> None:
    if not experiences:
        return

    for experience in experiences:
        _add_item_heading(document, experience.heading)
        _add_bullets(document, experience.bullets)


def _add_recruiter_projects(document: Document, projects: list[CvProject]) -> None:
    if not projects:
        return

    for project in projects:
        _add_item_heading(document, project.name)
        lines: list[str] = []
        if project.url:
            lines.append(f"URL: {project.url}")
        lines.extend(project.bullets)
        if project.technologies:
            lines.append("Technologies: " + _inline_list(project.technologies))
        _add_bullets(document, lines)


def _add_item_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(35, 35, 35)


def _add_experiences(document: Document, experiences: list[CvExperience]) -> None:
    if not experiences:
        _add_bullets(document, [], empty="Aucune expérience structurée détectée dans le profil maître.")
        return

    for experience in experiences:
        document.add_heading(experience.heading, level=2)
        _add_bullets(document, experience.bullets)


def _add_projects(document: Document, projects: list[CvProject], *, include_source: bool) -> None:
    if not projects:
        _add_bullets(document, [], empty="Aucun projet structuré détecté dans le profil maître.")
        return

    for project in projects:
        document.add_heading(project.name, level=2)
        lines = [f"Source: {project.source}"] if include_source else []
        if project.url:
            lines.append(f"URL: {project.url}")
        lines.extend(project.bullets)
        if project.technologies:
            lines.append("Technologies: " + _inline_list(project.technologies))
        _add_bullets(document, lines)


def _add_bullets(document: Document, items: list[str], *, empty: str | None = None) -> None:
    if not items:
        if empty is not None:
            document.add_paragraph(empty, style="List Bullet")
        return

    for item in items:
        paragraph = document.add_paragraph(item, style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(1)


def _safe_filename(cv_data: TargetedCvData, mode: CvRenderMode) -> str:
    company = _slug(cv_data.offer_info.company)
    title = _slug(cv_data.offer_info.title)
    prefix = "CV_Bastien" if mode == "recruiter" else "CV_Draft_Bastien"
    return f"{prefix}_{company}_{title}.docx"


def _slug(value: Any) -> str:
    text = unicodedata.normalize("NFKD", str(value).strip().lower())
    text = "".join(char for char in text if not unicodedata.combining(char))
    text = re.sub(r"[^a-z0-9]+", "_", text)
    text = re.sub(r"_+", "_", text).strip("_")
    return text[:80] or "non_renseigne"


def _inline_list(items: list[str]) -> str:
    return ", ".join(items) if items else "Aucune"


def _validate_mode(mode: str) -> None:
    if mode not in {"draft", "recruiter"}:
        raise ValueError("--mode doit etre 'draft' ou 'recruiter'.")
